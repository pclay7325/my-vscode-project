from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from database import SessionLocal, init_db, User
from models import ProductionPerformance
from utils.auth_utils import get_password_hash, verify_password, create_access_token, decode_access_token
import pandas as pd
from io import BytesIO
from datetime import timedelta
from fastapi.responses import StreamingResponse
from docx import Document  # For DOCX report generation

# Initialize FastAPI app
app = FastAPI(
    title="Manufacturing Analytics Backend",
    description="API for managing production performance data.",
    version="1.0.0"
)

# Initialize database
init_db()

# OAuth2 password flow for obtaining tokens
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/token")

# Dependency to get database session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --------------------- Pydantic Models ----------------------

class RegisterUserInput(BaseModel):
    username: str = Field(..., min_length=3, max_length=50, description="Username for registration")
    password: str = Field(..., min_length=6, max_length=50, description="Password with at least 6 characters")

# --------------------- AUTHENTICATION ENDPOINTS ----------------------

@app.post("/register", tags=["Auth"], summary="Register a new user")
def register_user(user_input: RegisterUserInput, db: Session = Depends(get_db)):
    """
    Register a new user by providing a **username** and **password**.
    """
    user = db.query(User).filter(User.username == user_input.username).first()
    if user:
        raise HTTPException(status_code=400, detail="Username already exists")
    hashed_password = get_password_hash(user_input.password)
    new_user = User(username=user_input.username, hashed_password=hashed_password)
    db.add(new_user)
    db.commit()
    return {"message": "User registered successfully"}

@app.post("/token", tags=["Auth"], summary="Login and get access token")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    """
    Login to obtain a JWT access token.
    """
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=30)
    access_token = create_access_token(data={"sub": user.username}, expires_delta=access_token_expires)
    return {"access_token": access_token, "token_type": "bearer"}

# --------------------- FILE UPLOAD ENDPOINT ----------------------

@app.post("/api/v1/upload/upload-production-performance", tags=["Upload"], summary="Upload production performance data")
async def upload_production_performance(
    file: UploadFile = File(...),
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):
    """
    Upload production performance data file (CSV or Excel).
    Supports flexible column names and requires a valid JWT token.
    """
    payload = decode_access_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

    # Flexible column name mapping
    column_mapping = {
        "machine": ["Machine", "machine"],
        "runtime_minutes": ["Runtime (Minutes)", "runtime_minutes"],
        "planned_time_minutes": ["Planned Time (Minutes)", "planned_time_minutes"],
        "good_units": ["Good Units", "good_units"],
        "total_units": ["Total Units", "total_units"],
        "downtime_minutes": ["Downtime (Minutes)", "downtime_minutes"]
    }

    def map_columns(df):
        """Map flexible column names to the required schema."""
        mapped_columns = {}
        for required, variations in column_mapping.items():
            for col in variations:
                if col in df.columns:
                    mapped_columns[required] = col
                    break
            else:
                raise HTTPException(status_code=400, detail=f"Missing required column: {required}")
        return df.rename(columns=mapped_columns)

    contents = await file.read()
    try:
        # Parse the uploaded file
        if file.filename.endswith('.csv'):
            df = pd.read_csv(BytesIO(contents))
        elif file.filename.endswith('.xlsx'):
            df = pd.read_excel(BytesIO(contents))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Upload CSV or Excel files.")

        # Map columns and validate data
        df = map_columns(df)
        for _, row in df.iterrows():
            new_record = ProductionPerformance(
                machine=row["machine"],
                runtime_minutes=row["runtime_minutes"],
                planned_time_minutes=row["planned_time_minutes"],
                good_units=row.get("good_units", 0),
                total_units=row.get("total_units", 0),
                downtime_minutes=row.get("downtime_minutes", 0),
            )
            db.add(new_record)
        db.commit()
        return {"message": "File uploaded and data processed successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")

# --------------------- REPORT GENERATION ENDPOINT ----------------------

@app.get("/api/v1/reports", tags=["Reports"], summary="Generate production performance report")
def generate_report(
    format: str = "json",
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):
    """
    Generate production performance report as JSON, CSV, Excel, or DOCX.

    **Available Formats**:
    - `json`: Returns the report as JSON.
    - `csv`: Downloads the report as a CSV file.
    - `excel`: Downloads the report as an Excel file.
    - `doc`: Downloads the report as a Word document (.docx).
    """
    payload = decode_access_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

    records = db.query(ProductionPerformance).all()
    if not records:
        raise HTTPException(status_code=404, detail="No production performance data available.")

    data = [
        {
            "machine": r.machine,
            "runtime_minutes": r.runtime_minutes,
            "planned_time_minutes": r.planned_time_minutes,
            "good_units": r.good_units,
            "total_units": r.total_units,
            "downtime_minutes": r.downtime_minutes,
        }
        for r in records
    ]

    df = pd.DataFrame(data)

    if format == "json":
        return data
    elif format == "csv":
        stream = BytesIO()
        df.to_csv(stream, index=False)
        stream.seek(0)
        return StreamingResponse(stream, media_type="text/csv", headers={"Content-Disposition": "attachment; filename=report.csv"})
    elif format == "excel":
        stream = BytesIO()
        with pd.ExcelWriter(stream, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Report")
        stream.seek(0)
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                 headers={"Content-Disposition": "attachment; filename=report.xlsx"})
    elif format == "doc":
        document = Document()
        document.add_heading("Production Performance Report", level=1)
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Add headers
        for i, column_name in enumerate(df.columns):
            table.rows[0].cells[i].text = column_name

        # Add data
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": "attachment; filename=report.docx"})
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'json', 'csv', 'excel', or 'doc'.")

# --------------------- ROOT ENDPOINT ----------------------

@app.get("/", tags=["Root"], summary="Welcome Endpoint")
def read_root():
    """
    Root endpoint with a welcome message.
    """
    return {"message": "Welcome to the Manufacturing Analytics Backend!"}
